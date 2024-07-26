from flask import Flask, request, jsonify
import pandas as pd
from PyPDF2 import PdfReader
import os
from docx import Document



app = Flask(__name__)
UPLOAD_FOLDER = r'C:\NIC internship work\project\uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Route to handle Excel file upload
@app.route('/upload/excel', methods=['POST'])
def upload_excel():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.xlsx'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        df = pd.read_excel(file_path)
        return jsonify(df.to_dict()), 200
    return jsonify({'error': 'Invalid file format'}), 400

# Route to handle PDF file upload
@app.route('/upload/pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.pdf'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
        return jsonify({'text': text}), 200
    return jsonify({'error': 'Invalid file format'}), 400

# Route to handle document file upload
@app.route('/upload/document', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and (file.filename.endswith('.txt') or file.filename.endswith('.docx')):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        if file.filename.endswith('.txt'):
            with open(file_path, 'r') as f:
                text = f.read()
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        return jsonify({'text': text}), 200
    return jsonify({'error': 'Invalid file format'}), 400

# New route to handle multiple file types
@app.route('/upload/all', methods=['POST'])
def upload_all():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    if file.filename.endswith('.xlsx'):
        df = pd.read_excel(file_path)
        return jsonify({'data': df.to_dict()}), 200
    elif file.filename.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
        return jsonify({'text': text}), 200
    elif file.filename.endswith('.txt'):
        with open(file_path, 'r') as f:
            text = f.read()
        return jsonify({'text': text}), 200
    elif file.filename.endswith('.docx'):
        doc = Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return jsonify({'text': text}), 200
    
    return jsonify({'error': 'Invalid file format'}), 400

if __name__ == '__main__':
    app.run(debug=True)
