from flask import Flask, request, jsonify
import pandas as pd
from PyPDF2 import PdfReader
import os
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer

app = Flask(__name__)
UPLOAD_FOLDER = "D:/Disruptive Ai/New folder"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

@app.route('/upload/excel', methods=['POST'])
def upload_excel():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.xlsx'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        try:
            df = pd.read_excel(file_path)
            text = df.to_string()
            vector = vectorize_text(text)
        finally:
            os.remove(file_path)
        return jsonify({'vector': vector.tolist()}), 200
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.pdf'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        text = ''
        try:
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
            vector = vectorize_text(text)
        finally:
            os.remove(file_path)
        return jsonify({'vector': vector.tolist()}), 200
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/document', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and (file.filename.endswith('.txt') or file.filename.endswith('.docx')):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        text = ''
        try:
            if file.filename.endswith('.txt'):
                with open(file_path, 'r') as f:
                    text = f.read()
            elif file.filename.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
            vector = vectorize_text(text)
        finally:
            os.remove(file_path)
        return jsonify({'vector': vector.tolist()}), 200
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/all', methods=['POST'])
def upload_all():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)
    
    text = ''
    try:
        if file.filename.endswith('.xlsx'):
            df = pd.read_excel(file_path)
            text = df.to_string()
        elif file.filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
        elif file.filename.endswith('.txt'):
            with open(file_path, 'r') as f:
                text = f.read()
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        if text:
            vector = vectorize_text(text)
        else:
            return jsonify({'error': 'No text found in file'}), 400
    finally:
        os.remove(file_path)
    
    return jsonify({'vector': vector.tolist()}), 200

def vectorize_text(text):
    try:
        vectorizer = TfidfVectorizer()
        vectors = vectorizer.fit_transform([text])
        return vectors.toarray()[0]
    except Exception as e:
        return {'error': str(e)}

if __name__ == '__main__':
    app.run(debug=True)
