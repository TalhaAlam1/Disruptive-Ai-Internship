from flask import Flask, request, jsonify
import pandas as pd
from PyPDF2 import PdfReader
import os
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
import firebase_admin
from firebase_admin import credentials, firestore
import re

app = Flask(__name__)
UPLOAD_FOLDER = "D:/Disruptive Ai/New folder"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize Firebase Admin SDK
cred = credentials.Certificate('D:/Disruptive Ai/agent-79934-firebase-adminsdk-ll7wr-2f2c46853e.json')  # Update this path
firebase_admin.initialize_app(cred)
db = firestore.client()  # Initialize Firestore client

def vectorize_text(text):
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([text])
    return vectors.toarray()[0].tolist()  # Convert NumPy array to list

def sanitize_id(filename):
    # Remove any characters not allowed in Firestore document IDs
    return re.sub(r'[^\w\s]', '_', filename)

@app.route('/upload/excel', methods=['POST'])
def upload_excel():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.xlsx'):
        try:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            df = pd.read_excel(file_path)
            text = df.to_string()
            vector = vectorize_text(text)
            doc_id = sanitize_id(f"excel_{file.filename}")
            # Store data in Firestore
            db.collection('documents').document(doc_id).set({
                'text': text,
                'vector': vector
            })
            return jsonify({'vector': vector}), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/pdf', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and file.filename.endswith('.pdf'):
        try:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            text = ''
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
            vector = vectorize_text(text)
            doc_id = sanitize_id(f"pdf_{file.filename}")
            # Store data in Firestore
            db.collection('documents').document(doc_id).set({
                'text': text,
                'vector': vector
            })
            return jsonify({'vector': vector}), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/document', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and (file.filename.endswith('.txt') or file.filename.endswith('.docx')):
        try:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            if file.filename.endswith('.txt'):
                with open(file_path, 'r') as f:
                    text = f.read()
            elif file.filename.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
            vector = vectorize_text(text)
            doc_id = sanitize_id(f"document_{file.filename}")
            # Store data in Firestore
            db.collection('documents').document(doc_id).set({
                'text': text,
                'vector': vector
            })
            return jsonify({'vector': vector}), 200
        except Exception as e:
            return jsonify({'error': str(e)}), 500
    return jsonify({'error': 'Invalid file format'}), 400

@app.route('/upload/all', methods=['POST'])
def upload_all():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        
        text = ''
        if file.filename.endswith('.xlsx'):
            df = pd.read_excel(file_path)
            text = df.to_string()
        elif file.filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text()
        elif file.filename.endswith('.txt'):
            with open(file_path, 'r') as f:
                text = f.read()
        elif file.filename.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
        
        if text:
            vector = vectorize_text(text)
            doc_id = sanitize_id(f"all_{file.filename}")
            # Store data in Firestore
            db.collection('documents').document(doc_id).set({
                'text': text,
                'vector': vector
            })
            return jsonify({'vector': vector}), 200
        return jsonify({'error': 'No text extracted'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
